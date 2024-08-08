import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Define the task data
tasks = [
    {"Task ID": "T1", "Task Description": "Define project goals", "Start Date": "25-Jul-2024", "End Date": "26-Jul-2024", "Duration (days)": 2},
    {"Task ID": "T2", "Task Description": "Identify stakeholders", "Start Date": "27-Jul-2024", "End Date": "28-Jul-2024", "Duration (days)": 2},
    {"Task ID": "T3", "Task Description": "Develop project charter", "Start Date": "29-Jul-2024", "End Date": "30-Jul-2024", "Duration (days)": 2},
    {"Task ID": "T4", "Task Description": "Conduct market research", "Start Date": "31-Jul-2024", "End Date": "2-Aug-2024", "Duration (days)": 3},
    {"Task ID": "T5", "Task Description": "Analyze target audience needs", "Start Date": "3-Aug-2024", "End Date": "5-Aug-2024", "Duration (days)": 3},
    {"Task ID": "T6", "Task Description": "Identify technology requirements", "Start Date": "6-Aug-2024", "End Date": "7-Aug-2024", "Duration (days)": 2},
    {"Task ID": "T7", "Task Description": "Design user interface", "Start Date": "8-Aug-2024", "End Date": "9-Aug-2024", "Duration (days)": 2},
    {"Task ID": "T8", "Task Description": "Develop front-end", "Start Date": "10-Aug-2024", "End Date": "12-Aug-2024", "Duration (days)": 3},
    {"Task ID": "T9", "Task Description": "Develop back-end", "Start Date": "10-Aug-2024", "End Date": "13-Aug-2024", "Duration (days)": 4},
    {"Task ID": "T10", "Task Description": "Integrate features", "Start Date": "14-Aug-2024", "End Date": "15-Aug-2024", "Duration (days)": 2},
    {"Task ID": "T11", "Task Description": "Test web application", "Start Date": "16-Aug-2024", "End Date": "17-Aug-2024", "Duration (days)": 2},
    {"Task ID": "T12", "Task Description": "Launch web application", "Start Date": "18-Aug-2024", "End Date": "18-Aug-2024", "Duration (days)": 1}
]

# Convert to DataFrame for easy manipulation
df_tasks = pd.DataFrame(tasks)

# Create a Word document
doc = Document()

# Add a table with rounded edges
table = doc.add_table(rows=1, cols=len(df_tasks.columns))
table.style = 'Table Grid'

# Add header
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(df_tasks.columns):
    hdr_cells[i].text = column_name
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.font.size = Pt(12)
    hdr_cells[i].paragraphs[0].alignment = 1  # Center alignment

# Add body rows
for index, row in df_tasks.iterrows():
    row_cells = table.add_row().cells
    for i, cell in enumerate(row):
        row_cells[i].text = str(cell)
        run = row_cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(12)
        row_cells[i].paragraphs[0].alignment = 1  # Center alignment

# Function to set table borders
def set_table_borders(table):
    tbl = table._element
    tbl_pr = tbl.find('.//w:tblPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    
    # Remove existing border elements
    tbl_borders = tbl_pr.find('.//w:tblBorders', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if tbl_borders is not None:
        tbl_pr.remove(tbl_borders)
    
    # Add new border elements
    border_elm = OxmlElement("w:tblBorders")
    border_elm.set(qn("w:top"), "nil")
    border_elm.set(qn("w:left"), "nil")
    border_elm.set(qn("w:bottom"), "nil")
    border_elm.set(qn("w:right"), "nil")
    border_elm.set(qn("w:insideH"), "nil")
    border_elm.set(qn("w:insideV"), "nil")
    tbl_pr.append(border_elm)

# Apply borders to the table
set_table_borders(table)

# Save the document with a proper file name
doc.save("tasks_document.docx")
